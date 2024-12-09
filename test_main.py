import os
import unittest
import shutil
from docx import Document
from main import replace_text_within_percent_signs

class TestReplaceText(unittest.TestCase):
    def setUp(self):
        self.tmp_folder_name = "temporary_folder"
        self.tmp_folder_path = os.path.join(os.getcwd(), f"{self.tmp_folder_name}_TMP")
        shutil.rmtree(self.tmp_folder_path, ignore_errors=True)
        os.makedirs(self.tmp_folder_path)

    def tearDown(self):
        shutil.rmtree(self.tmp_folder_path)

    def test_replace_text_within_percent_signs(self):
        replace_dict = {
            "標案案號": "YL113501",
        }
        file_path = "test.docx"

        # 創建一個測試文件
        doc = Document()
        doc.add_paragraph("這是一個%%標案案號%%的測試。")
        doc.save(file_path)

        # 調用函數進行替換
        replace_text_within_percent_signs(file_path, replace_dict, self.tmp_folder_name)

        # 驗證結果
        doc = Document(os.path.join(self.tmp_folder_path, "test.docx"))
        for para in doc.paragraphs:
            self.assertIn("這是一個YL113501的測試。", para.text)

    def test_replace_text_within_percent_signs_multiple_replacements(self):
        replace_dict = {
            "標案案號": "YL113501",
            "標案名稱": "測試標案",
        }
        file_path = "test.docx"

        # 創建一個測試文件
        doc = Document()
        doc.add_paragraph("這是一個%%標案案號%%的%%標案名稱%%。")
        doc.save(file_path)

        # 調用函數進行替換
        replace_text_within_percent_signs(file_path, replace_dict, self.tmp_folder_name)

        # 驗證結果
        doc = Document(os.path.join(self.tmp_folder_path, "test.docx"))
        for para in doc.paragraphs:
            self.assertIn("這是一個YL113501的測試標案。", para.text)

    def test_replace_text_within_percent_signs_no_replacement(self):
        replace_dict = {
            "標案案號": "YL113501",
        }
        file_path = "test.docx"

        # 創建一個測試文件
        doc = Document()
        doc.add_paragraph("這是一個測試。")
        doc.save(file_path)

        # 調用函數進行替換
        replace_text_within_percent_signs(file_path, replace_dict, self.tmp_folder_name)

        # 驗證結果
        doc = Document(os.path.join(self.tmp_folder_path, "test.docx"))
        for para in doc.paragraphs:
            self.assertIn("這是一個測試。", para.text)

if __name__ == '__main__':
    unittest.main()
