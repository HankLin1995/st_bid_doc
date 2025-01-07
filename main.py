import os
import shutil
from docx import Document
# from docx.enum.text import WD_COLOR_INDEX

def replace_text_within_percent_signs(file_path, replace_dict):#, tmp_folder_name):
    # tmp_folder_path = os.path.join(os.getcwd(), f"{tmp_folder_name}_TMP")
    
    # if not os.path.exists(tmp_folder_path):
        # os.makedirs(tmp_folder_path)
    
    # file_name = os.path.basename(file_path)
    # tmp_file_path = os.path.join(tmp_folder_path, file_name)
    
    # shutil.copy(file_path, tmp_file_path)
    
    try:
        doc = Document(file_path)
        
        for para in doc.paragraphs:
            replace_in_paragraph(para, replace_dict)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, replace_dict)

        doc.save(file_path)
    
    except Exception as e:
        print(f"無法處理文件 {file_path}: {e}")
    
    #shutil.rmtree(tmp_folder_path)  

def replace_in_paragraph(paragraph, replace_dict):
    """替換段落中的標記文字，保留原始格式。"""
    # 獲取原始的 runs
    original_runs = paragraph.runs
    
    # 重新組合段落中的文字
    text = paragraph.text
    start_pos = text.find("%%")

    while start_pos != -1:
        end_pos = text.find("%%", start_pos + 2)
        
        if end_pos != -1:
            target_content = text[start_pos + 2:end_pos]
            
            if target_content in replace_dict:
                # 替換的文字
                replace_with = replace_dict[target_content]

                # 遍歷段落中的 runs 找到包含標記的 run 範圍
                current_pos = 0
                start_run_index = None
                end_run_index = None
                
                for i, run in enumerate(paragraph.runs):
                    next_pos = current_pos + len(run.text)
                    if current_pos <= start_pos < next_pos:
                        start_run_index = i
                    if current_pos < end_pos + 2 <= next_pos:
                        end_run_index = i
                        break
                    current_pos = next_pos

                # 確保找到了對應的 run
                if start_run_index is not None and end_run_index is not None:
                    # 保存標記的文字格式
                    first_run = paragraph.runs[start_run_index]
                    
                    # 刪除標記所在的 runs
                    for i in range(end_run_index, start_run_index - 1, -1):
                        paragraph._p.remove(paragraph.runs[i]._element)
                    
                    # 插入替換的文字
                    new_run = paragraph.add_run(replace_with)
                    
                    # 保持新文字的格式
                    new_run.font.name = first_run.font.name
                    new_run.font.size = first_run.font.size
                    new_run.font.bold = first_run.font.bold
                    new_run.font.italic = first_run.font.italic
                    new_run.font.underline = first_run.font.underline
                    if hasattr(first_run.font, 'color'):
                        new_run.font.color.rgb = first_run.font.color.rgb

                    print(f"已將 %%{target_content}%% 替換為 {replace_with}")
                    
                    # 更新段落文本
                    text = paragraph.text
                    start_pos = text.find("%%")
                    continue

        start_pos = text.find("%%", start_pos + 2)

