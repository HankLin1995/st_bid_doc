
from docx import Document
from docx.shared import RGBColor

def replace_text_within_percent_signs(file_path, replace_dict):
    
    try:
        doc = Document(file_path)
        
        for para in doc.paragraphs:
            replace_in_paragraph(para, replace_dict)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, replace_dict)

        doc.save(file_path)

        if check_replaced(file_path, replace_dict)==False:
            return f"文件 {file_path} 有未處理的標記。"
    
    except Exception as e:
        print(f"無法處理文件 {file_path}: {e}")
    
    #shutil.rmtree(tmp_folder_path)  

def replace_in_paragraph(paragraph, replace_dict):
    """基於整段文字進行標記替換，並保留格式。"""
    text = paragraph.text  # 獲取整段文字
    replaced_text = text  # 用於存放替換後的文字

    # 找到所有標記並替換
    for key, value in replace_dict.items():
        replaced_text = replaced_text.replace(f"%%{key}%%", str(value))

    # 如果沒有任何替換，直接返回
    if replaced_text == text:
        return

    # 清空原始 runs
    for run in paragraph.runs:
        paragraph._p.remove(run._element)

    # 將替換後的文字重新加入 paragraph，並保留第一個 run 的格式
    first_run = paragraph.add_run(replaced_text)

    # 保留第一個 run 的格式
    if paragraph.runs:
        original_run = paragraph.runs[0]
        first_run.font.name = original_run.font.name
        first_run.font.size = original_run.font.size
        first_run.font.bold = original_run.font.bold
        first_run.font.italic = original_run.font.italic
        first_run.font.underline = original_run.font.underline
        if hasattr(original_run.font, 'color'):
            first_run.font.color.rgb = original_run.font.color.rgb
        # Set the font color to red
        first_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
# check if the file content is all replaced (contains no %%)

def check_replaced(file_path, replace_dict):
    doc = Document(file_path)
    for para in doc.paragraphs:
        text = para.text
        if "%%" in text:
            return False
        for key in replace_dict.keys():
            if f"%%{key}%%" in text:
                return False

    # print(f"已完成段落替換: {replaced_text}")

#------------------------------------------------------------------------------#
# deal with di

import xml.etree.ElementTree as ET
import io
import re

def read_tender_document(file_path, replacements=None, output_file=None):
    try:
        if replacements is None:
            replacements = {}
            
        # 讀取XML文件內容
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            
        # 找出所有需要替換的項目
        pattern = r'%%([^%]+)%%'
        required_replacements = set(re.findall(pattern, content))
        provided_replacements = set(replacements.keys())
        
        # 檢查是否有遺漏的替換項目
        missing_replacements = required_replacements - provided_replacements
        if missing_replacements:
            print("\n警告：以下項目尚未提供替換值：")
            for item in missing_replacements:
                print(f"- {item}")
            # return
            
        # # 檢查是否有多餘的替換項目
        # extra_replacements = provided_replacements - required_replacements
        # if extra_replacements:
        #     print("\n注意：以下替換項目在文件中未被使用：")
        #     for item in extra_replacements:
        #         print(f"- {item}")
            
        # 替換所有%%包圍的文字
        for key, value in replacements.items():
            pattern = f'%%{key}%%'
            content = content.replace(pattern, value)
            
        # 最後檢查是否還有任何未替換的%%標記
        remaining_patterns = re.findall(r'%%[^%]+%%', content)
        if remaining_patterns:
            print("\n錯誤：仍有未被替換的項目：")
            for pattern in remaining_patterns:
                print(f"- {pattern}")
            # return 
            
        # 如果指定了輸出文件，將處理後的內容保存到新文件
        if output_file:
            with open(output_file, 'w', encoding='utf-8') as file:
                file.write(content)
            print(f"\n處理後的文件已保存至: {output_file}")
            
    #     # 將修改後的內容寫入臨時字符串以供 ET 解析
    #     tree = ET.parse(io.StringIO(content))
    #     root = tree.getroot()
        
    #     # 獲取基本信息
    #     unit = root.find('.//單位名').text
    #     subject = root.find('.//主旨/文字').text
        
    #     # 獲取所有說明段落
    #     explanations = root.findall('.//段落[@段名="說明："]/條列')
        
    #     print(f"單位: {unit}")
    #     print(f"主旨: {subject}")
    #     print("\n說明事項:")
        
    #     # 打印所有說明段落
    #     for item in explanations:
    #         number = item.get('序號', '')
    #         text = ''.join(item.itertext()).strip()
    #         print(f"{number} {text}")
            
    #     # 獲取署名和日期
    #     signature = root.find('.//署名').text
    #     date = root.find('.//年月日').text
        
    #     print(f"\n署名: {signature}")
    #     print(f"日期: {date}")
        
    except Exception as e:
        print(f"錯誤：{str(e)}")

# if __name__ == "__main__":
#     file_path = "預算書簽(新)-工程-未核定-委外.txt"
#     output_file = "招標簽_已處理.txt"  # 新的輸出文件名
    
#     # 定義要替換的文字
#     replacements = {
#         "工程名稱": "2024年度灌溉渠道整建工程",
#         "民國年":"113",
#         "所屬分處": "台北分處",
#         "委外廠商":"OOO有限公司",
#         "工程編號": "113-A-01-01",
#         "經費來源": "112年度資本支出預算",
#         "採購標的": "工程",
#         "核定經費": "新台幣1,500,000元整",
#         "預算書總價": "新台幣1,450,000元整",
#         "發包工作費總額": "新台幣1,400,000元整",
#         "工程分類":"未達二千萬之第三類工程",#二千萬元以上未達查核金額之第二類工程
#         # 新增缺少的項目
#         "押標金額度": "新台幣70,000元整",
#         "廠商基本資格": "土木包工業或丙等以上綜合營造業",
#         "採購金額級距": "未達公告金額之採購",
#         "履約保證金": "新台幣140,000元整"
#     }
    
#     read_tender_document(file_path, replacements, output_file)


